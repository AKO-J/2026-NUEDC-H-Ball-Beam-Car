from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("智能车_正方形循迹_最新状态.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "5B6470"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
LATIN_FONT = "Calibri"
# STHeiti is bundled with macOS and renders Chinese reliably in the local
# LibreOffice-based DOCX preview.  Set it explicitly as the East Asian font;
# otherwise Chinese characters may fall back to missing-glyph boxes.
CJK_FONT = "Hiragino Sans GB"


def apply_fonts(r_pr):
    fonts = r_pr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)


def set_run_font(run, size, color=INK, bold=False, italic=False):
    run.font.name = LATIN_FONT
    apply_fonts(run._element.get_or_add_rPr())
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table_header(row, values):
    for cell, value in zip(row.cells, values):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, 10, DARK_BLUE, bold=True)


def add_table_row(row, values, emphasis_first=False):
    for index, (cell, value) in enumerate(zip(row.cells, values)):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, 10, INK, bold=(emphasis_first and index == 0))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}  ")
    set_run_font(r, 10.5, DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, 10.5, INK)
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("智能车循迹 · 当前状态 · 2026-07-26")
    set_run_font(run, 8.5, MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_SECTION.NEW_PAGE if False else section.orientation
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    apply_fonts(normal._element.get_or_add_rPr())
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = LATIN_FONT
        apply_fonts(style._element.get_or_add_rPr())
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        apply_fonts(style._element.get_or_add_rPr())
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_document():
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("智能车 · 当前技术状态")
    set_run_font(r, 24, INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("正方形循迹与 IMU 90° 转弯 · 最新精简交接文档")
    set_run_font(r, 12, MUTED)

    meta = doc.add_table(rows=2, cols=2)
    set_table_geometry(meta, [1800, 7560])
    labels = [("更新时间", "2026-07-26"), ("当前固件", "已构建并烧录：build/line_follow_test.out")]
    for row, (label, value) in zip(meta.rows, labels):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        add_table_row(row, (label, value), emphasis_first=True)

    doc.add_paragraph()
    add_callout(
        doc,
        "当前结论",
        "软件已具备“角点触发 → 主动刹车 → ICM42688 累计 90° → 重新找线”的完整流程。"
        "是否能稳定通过正方形赛道尚未完成地面验收，下一步必须以单个直角实测结果继续校准。",
    )

    add_heading(doc, "1. 已确认的硬件与标定", 1)
    hardware = doc.add_table(rows=1, cols=3)
    set_table_geometry(hardware, [2500, 2700, 4160])
    add_table_header(hardware.rows[0], ("项目", "当前确认", "备注"))
    rows = [
        ("主控与驱动", "MSPM0G3507 + TB6612", "电机 A 为物理左轮，B 为物理右轮。"),
        ("灰度模块", "8 路 Yahboom 模块", "车头向前时 X1 在左；黑线原始电平为 0。"),
        ("IMU", "ICM42688", "I2C 通信已实测成功；Z 轴用于航向角累计。"),
        ("几何参数", "轮距 112 mm；传感器前伸 152 mm", "X1–X8 跨度 79 mm；离地 19 mm。"),
        ("赛道参数", "黑胶带宽 19 mm；最小弯半径 400 mm", "当前目标是稳定通过 90° 方形角。"),
    ]
    for values in rows:
        row = hardware.add_row()
        add_table_row(row, values, emphasis_first=True)

    add_heading(doc, "2. 当前有效控制逻辑", 1)
    add_callout(doc, "角点启动条件", "先沿窄线（≤3 路黑）运行 24 个控制周期，约 96 ms；随后同侧出现 ≥4 路黑的窄→宽跃升才认定为角点。静态放在宽黑区域不会立即转弯。")

    states = doc.add_table(rows=1, cols=3)
    set_table_geometry(states, [1500, 2600, 5260])
    add_table_header(states.rows[0], ("状态", "核心动作", "退出条件"))
    rows = [
        ("0 FOLLOW", "4 ms 控制周期，单次 8 路扫描。", "检测到有效窄→宽角点边沿。"),
        ("1 BRAKE", "TB6612 双轮主动短刹 24 ms。", "进入陀螺仪转弯。"),
        ("2 TURN", "单轮 PIVOT；IMU 连续累计航向。", "累计达到 90,000 mdeg。"),
        ("3 REACQUIRE", "18% 低速前进找新边。", "连续 2 帧检测到窄线；最多 600 ms。"),
        ("4 FAULT", "安全停车。", "IMU 读错、转弯超时或找线超时。"),
    ]
    for values in rows:
        row = states.add_row()
        add_table_row(row, values, emphasis_first=True)

    add_bullet(doc, "右侧支路（X5–X8）：左轮驱动、右轮 0%。")
    add_bullet(doc, "左侧支路（X1–X4）：左轮 0%、右轮驱动。")
    add_bullet(doc, "转弯最后阶段从 72° 开始限制为低速；IMU 到 90° 后再主动短刹 32 ms。")

    add_heading(doc, "3. 当前关键参数", 1)
    params = doc.add_table(rows=1, cols=3)
    set_table_geometry(params, [3300, 1500, 4560])
    add_table_header(params.rows[0], ("参数", "当前值", "用途"))
    rows = [
        ("LINE_FOLLOW_CONTROL_SLICE_MS", "4 ms", "普通循迹和角点读取周期。"),
        ("SQUARE_PIVOT_ARM_NARROW_SLICES", "24", "约 96 ms 窄线预置，防止静态误转。"),
        ("SQUARE_BRAKE_MS", "24 ms", "角点触发后的主动短刹。"),
        ("SQUARE_TURN_TARGET_MDEG", "90,000", "陀螺仪 90° 转弯结束条件。"),
        ("SQUARE_TURN_SLOWDOWN_MDEG / DUTY", "72,000 / 18%", "最后 18° 的低速收角。"),
        ("右侧支路左轮 PWM", "41% → 18%", "右转时的主驱动轮。"),
        ("左侧支路右轮 PWM", "20% 起转 → 10% 保持", "左转时的主驱动轮。"),
        ("SQUARE_TURN_FINISH_BRAKE_MS", "32 ms", "IMU 到 90° 后抑制惯性。"),
        ("SQUARE_REACQUIRE_DUTY / MAX_MS", "18% / 600 ms", "转完后重新找到下一条线。"),
    ]
    for values in rows:
        row = params.add_row()
        add_table_row(row, values, emphasis_first=True)

    add_heading(doc, "4. 代码、构建与烧录", 1)
    add_bullet(doc, "主程序：motor_forward_test/line_follow_test.c")
    add_bullet(doc, "电机控制：motor_forward_test/motor_driver.c；其中 Motor_brakeAllFor() 为最新主动刹车接口。")
    add_bullet(doc, "灰度与 IMU：gray_sensor.c、icm42688.c；当前主程序使用单次 8 路扫描以提高角点采样速度。")
    add_bullet(doc, "构建：在 motor_forward_test 目录执行 make line-follow。")
    add_bullet(doc, "烧录产物：build/line_follow_test.out。烧录后上电前约 4 秒不得移动小车，以完成 IMU 零偏标定。")

    add_heading(doc, "5. 下一步地面验证（按顺序）", 1)
    steps = [
        "只布置“直线 + 一个右侧支路 90° 角”，从普通窄线段开始，不要把车直接放在宽黑区域。",
        "上电后保持静止约 4 秒；确认直线能正常前进且不会静态误转。",
        "观察角点序列：主动短刹 → 左轮驱动、右轮停 → 接近 90° 时降速并主动短刹 → 低速找新边。",
        "若角度不足/超过：优先改 SQUARE_TURN_TARGET_MDEG；若转后找不到线：改 SQUARE_REACQUIRE_DUTY。每次只调一项。",
        "连续通过单角 5 次后，再测试完整正方形赛道。",
    ]
    for item in steps:
        add_number(doc, item)

    add_heading(doc, "6. 调试观察值", 1)
    watch = doc.add_table(rows=1, cols=2)
    set_table_geometry(watch, [3300, 6060])
    add_table_header(watch.rows[0], ("CCS Watch 变量", "应看到的结果"))
    rows = [
        ("g_square_state", "0 → 1 → 2 → 3 → 0；若变为 4，代表安全故障停车。"),
        ("g_square_turn_angle_mdeg", "转弯结束应接近 90,000。"),
        ("g_square_turn_driven_wheel_is_left", "右侧支路应为 1；左侧支路应为 0。"),
        ("g_line_applied_left_duty / right_duty", "TURN 时可核对实际左右 PWM 是否符合物理轮命令。"),
        ("g_square_imu_error / g_square_turn_timed_out", "正常均应为 0。"),
    ]
    for values in rows:
        row = watch.add_row()
        add_table_row(row, values, emphasis_first=True)

    add_heading(doc, "7. 已废弃：不要恢复", 1)
    add_bullet(doc, "不再使用“黑线数降到一半以下即结束转弯”；它会在直角处过早结束。")
    add_bullet(doc, "不再使用内外轮速度比、编码器弯道比 2.17 或 SQUARE_TURN_REVERSE_STEERING 的反转映射。")
    add_bullet(doc, "不再把右轮设置为持续 1% PWM；该占空比通常无法使电机从静止起转。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
