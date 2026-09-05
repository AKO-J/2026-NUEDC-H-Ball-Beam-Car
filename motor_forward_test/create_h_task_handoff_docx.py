#!/usr/bin/env python3
"""Create a polished Word version of H_TASK_AI_HANDOFF.md.

Run with the workspace bundled Python.  This intentionally has no external
Markdown dependency so the handoff remains reproducible on the team Mac.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(__file__).with_name("H_TASK_AI_HANDOFF.md")
OUTPUT = Path(__file__).with_name("H_TASK_AI_HANDOFF_20260729.docx")

# Arial Unicode MS is copied into the isolated LibreOffice QA profile before
# rendering, so the generated DOCX keeps one portable, CJK-capable typeface.
FONT = "Arial Unicode MS"
MONO_FONT = "Menlo"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
TABLE_HEADER = "E8EEF5"
TABLE_BORDER = "B8C4D1"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF7E6"
RED = "9B1C1C"
USABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths, header=True):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "4")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), TABLE_BORDER)
        borders.append(item)
    tbl_pr.append(borders)
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        if index < len(grid.gridCol_lst):
            grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_index == 0:
                set_cell_shading(cell, TABLE_HEADER)
                tr_pr = row._tr.get_or_add_trPr()
                if tr_pr.find(qn("w:tblHeader")) is None:
                    tr_pr.append(OxmlElement("w:tblHeader"))


def add_inline(paragraph, text, size=11, color="000000", bold=False,
               italic=False):
    """Render bold, italic, and backtick spans without losing Chinese text."""
    pieces = re.split(r"(`[^`]*`|\*\*[^*]*\*\*|\*[^*]+\*)", text)
    for piece in pieces:
        if not piece:
            continue
        is_code = piece.startswith("`") and piece.endswith("`")
        is_bold = piece.startswith("**") and piece.endswith("**")
        is_italic = (piece.startswith("*") and piece.endswith("*") and
                     not is_bold)
        value = piece[1:-1] if is_code else (piece[2:-2] if is_bold else
                (piece[1:-1] if is_italic else piece))
        run = paragraph.add_run(value)
        set_run_font(run, MONO_FONT if is_code else FONT, size,
                     DARK_BLUE if is_code else color,
                     bold or is_bold, italic or is_italic)


def add_body_paragraph(doc, text, style_name="Normal", left_indent=None):
    p = doc.add_paragraph(style=style_name)
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)
    add_inline(p, text)
    return p


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    configure_table(table, [USABLE_WIDTH_DXA], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 2, 2, 1.2)
    add_inline(p, text.lstrip("> "), size=10.5, color=INK, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    configure_table(table, [USABLE_WIDTH_DXA], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 2, 2, 1.05)
    run = p.add_run("\n".join(lines))
    set_run_font(run, MONO_FONT, 8.8, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [part.strip() for part in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    return rows


def column_widths(rows):
    columns = max(len(row) for row in rows)
    lengths = [max(8, max(len(row[index]) if index < len(row) else 0
                          for row in rows)) for index in range(columns)]
    total = sum(lengths)
    widths = [max(950, int(USABLE_WIDTH_DXA * length / total))
              for length in lengths]
    difference = USABLE_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    return widths


def add_markdown_table(doc, markdown_lines):
    rows = parse_table(markdown_lines)
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    widths = column_widths(rows)
    for row_index, row in enumerate(rows):
        for col_index in range(columns):
            text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, 0, 0, 1.12)
            add_inline(p, text, size=9.6,
                       color=INK if row_index == 0 else "000000",
                       bold=(row_index == 0))
    configure_table(table, widths, header=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, 0, 0, 1)
    run = p.add_run("2026 电赛 H 题 · AI 交接与执行清单")
    set_run_font(run, FONT, 8.5, MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, 0, 0, 1)
    run = p.add_run("第 ")
    set_run_font(run, FONT, 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)
    run = p.add_run(" 页")
    set_run_font(run, FONT, 8.5, MUTED)


def add_cover(doc, title):
    for _ in range(5):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 0, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 8, 1.1)
    run = p.add_run(title)
    set_run_font(run, FONT, 25, INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 6, 1.1)
    run = p.add_run("当前状态、硬件边界与三日执行清单")
    set_run_font(run, FONT, 14, DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, 0, 0, 1.1)
    run = p.add_run("最后核对：2026-07-29")
    set_run_font(run, FONT, 10.5, MUTED)
    doc.add_page_break()


def markdown_to_docx(source: Path, output: Path):
    doc = Document()
    style_document(doc)
    lines = source.read_text(encoding="utf-8").splitlines()
    first_title_done = False
    index = 0
    in_code = False
    code_lines = []

    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            if level == 1 and not first_title_done:
                add_cover(doc, title)
                first_title_done = True
            else:
                p = doc.add_paragraph(style=f"Heading {level}")
                add_inline(p, title, size={1: 16, 2: 13, 3: 12}[level],
                           color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
                           bold=True)
            index += 1
            continue
        if line.startswith(">"):
            add_callout(doc, line)
            index += 1
            continue
        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        numbered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if bullet or numbered:
            match = bullet or numbered
            indent = len(match.group(1).expandtabs(2)) // 2
            style_name = "List Bullet" if bullet else "List Number"
            p = doc.add_paragraph(style=style_name)
            if indent:
                p.paragraph_format.left_indent = Inches(0.5 + 0.25 * indent)
            add_inline(p, match.group(2))
            index += 1
            continue
        if re.fullmatch(r"-{3,}", line.strip()):
            index += 1
            continue
        add_body_paragraph(doc, line)
        index += 1

    if in_code:
        add_code_block(doc, code_lines)
    doc.core_properties.title = "2026 电赛 H 题：AI 交接、当前状态与三日执行清单"
    doc.core_properties.subject = "H 题项目交接与执行约束"
    doc.core_properties.author = "SMARTCAR Team"
    doc.save(output)


if __name__ == "__main__":
    markdown_to_docx(SOURCE, OUTPUT)
    print(OUTPUT)
