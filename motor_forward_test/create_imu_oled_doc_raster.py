"""Build a Chinese-safe DOCX field guide for the current smart-car IMU/OLED setup.

LibreOffice in the local headless runtime cannot render macOS CJK run fonts
reliably.  The pages are therefore laid out with the installed Chinese font
first and placed into the DOCX as high-resolution pages.  This keeps the
delivered Word document readable in both Word and the local render check.
"""

from pathlib import Path
from tempfile import TemporaryDirectory

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "智能车_陀螺仪校准与OLED使用指南_2026-07-27.docx"

PAGE_W, PAGE_H = 1654, 2339  # A4 ratio, 200 dpi when inserted at A4 size
M = 112
FONT_FILE = "/System/Library/Fonts/Supplemental/Songti.ttc"
SANS_FILE = "/System/Library/Fonts/STHeiti Medium.ttc"

NAVY = "#112B46"
BLUE = "#1666A8"
TEAL = "#0B8574"
INK = "#183247"
MUTED = "#627080"
PALE = "#EDF4F8"
PALE_BLUE = "#E5F1FA"
PALE_GREEN = "#EAF6F1"
LINE = "#C9D7E1"
WHITE = "#FFFFFF"


def font(size, bold=False):
    # Songti index 0 is a CJK-capable face.  Heiti gives compact bold labels.
    if bold:
        return ImageFont.truetype(SANS_FILE, size, index=1)
    return ImageFont.truetype(FONT_FILE, size, index=0)


def fit_lines(draw, text, fnt, max_width):
    lines, line = [], ""
    for char in text:
        if char == "\n":
            lines.append(line)
            line = ""
            continue
        trial = line + char
        if line and draw.textlength(trial, font=fnt) > max_width:
            lines.append(line)
            line = char
        else:
            line = trial
    if line:
        lines.append(line)
    return lines or [""]


def text_block(draw, xy, text, fnt, fill=INK, width=None, leading=10):
    x, y = xy
    if width is None:
        lines = text.split("\n")
    else:
        lines = fit_lines(draw, text, fnt, width)
    bbox = draw.textbbox((0, 0), "国Ag", font=fnt)
    line_h = bbox[3] - bbox[1] + leading
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h
    return y


def box(draw, xy, fill=PALE, outline=None, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline)


def page_base(page_no, kicker):
    img = Image.new("RGB", (PAGE_W, PAGE_H), WHITE)
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, PAGE_W, 26), fill=TEAL)
    draw.text((M, 64), kicker, font=font(23, True), fill=BLUE)
    draw.line((M, 106, PAGE_W - M, 106), fill=LINE, width=2)
    draw.line((M, PAGE_H - 91, PAGE_W - M, PAGE_H - 91), fill=LINE, width=2)
    draw.text((M, PAGE_H - 67), "智能车 H 题｜现场调试参考｜2026-07-27", font=font(20), fill=MUTED)
    foot = f"{page_no} / 3"
    draw.text((PAGE_W - M - draw.textlength(foot, font=font(20)), PAGE_H - 67),
              foot, font=font(20), fill=MUTED)
    return img, draw


def heading(draw, y, text, n):
    draw.text((M, y), f"{n}  {text}", font=font(40, True), fill=NAVY)
    return y + 65


def bullet_list(draw, y, items, width=PAGE_W - 2 * M, size=28, gap=13):
    fnt = font(size)
    bullet = font(size, True)
    for item in items:
        draw.text((M + 6, y), "•", font=bullet, fill=TEAL)
        y = text_block(draw, (M + 39, y), item, fnt, width=width - 39, leading=8)
        y += gap
    return y


def simple_table(draw, y, headers, rows, widths, row_height=63, text_size=23):
    x = M
    hfont = font(text_size, True)
    bfont = font(text_size)
    for value, width in zip(headers, widths):
        draw.rectangle((x, y, x + width, y + row_height), fill=PALE_BLUE, outline=LINE)
        draw.text((x + 12, y + 14), value, font=hfont, fill=NAVY)
        x += width
    y += row_height
    for row in rows:
        x = M
        for value, width in zip(row, widths):
            draw.rectangle((x, y, x + width, y + row_height), fill=WHITE, outline=LINE)
            draw.text((x + 12, y + 14), value, font=bfont, fill=INK)
            x += width
        y += row_height
    return y


def create_pages(tmp):
    paths = []

    # Page 1: gyro wiring, calibration and validation.
    img, d = page_base(1, "SMART CAR / IMU CALIBRATION")
    d.text((M, 150), "陀螺仪校准与 OLED 状态显示", font=font(62, True), fill=NAVY)
    d.text((M, 237), "当前实现、接线、验证结果与现场操作", font=font(29), fill=MUTED)
    box(d, (M, 305, PAGE_W - M, 507), PALE_GREEN)
    d.text((M + 28, 332), "当前结论", font=font(30, True), fill=TEAL)
    text_block(d, (M + 28, 380),
               "最新测试：原地完整转一圈后 HDG = 356.6°，回到初始方向为 1.8°。\n"
               "航向标定可维持比例 1.00；重点是保持开机静止完成零偏采样。",
               font(28), width=PAGE_W - 2 * M - 56, leading=9)

    y = heading(d, 550, "陀螺仪接线（ICM-42688）", "1")
    y = simple_table(d, y, ["模块引脚", "TI 板连接", "说明"], [
        ("VCC", "3.3 V", "禁止接 5 V"),
        ("GND", "GND", "必须与电机系统共地"),
        ("SDA", "PA28 / I²C0 SDA / PINCM3", "姿态传感器数据线"),
        ("SCL", "PA31 / I²C0 SCL / PINCM6", "姿态传感器时钟线"),
        ("CS / AD0", "CS→3.3 V；AD0→GND", "I²C 地址为 0x68"),
    ], [222, 536, 560])

    y += 40
    y = heading(d, y, "当前零漂与积分参数", "2")
    y = simple_table(d, y, ["参数", "当前值", "用途"], [
        ("静止等待", "1 000 ms", "开机后先等待机械振动消失"),
        ("零偏采样", "256 次 × 5 ms", "只使用 Z 轴陀螺仪，得到初始 bias"),
        ("静止阈值", "±120 raw（约 ±7.3°/s）", "静止时不积分，同时缓慢跟踪零偏"),
        ("积分时间", "SysTick 实际经过时间", "不再假设每次循环固定 10 ms"),
        ("比例", "1.00（100 / 100）", "由本轮 356.6° 验证，当前不改"),
        ("零偏精度", "Q8 小数累积", "可吸收小于 1 raw 的持续漂移"),
    ], [234, 408, 676], row_height=66, text_size=22)

    y += 34
    box(d, (M, y, PAGE_W - M, y + 225), PALE)
    d.text((M + 24, y + 20), "启动步骤", font=font(29, True), fill=BLUE)
    text_block(d, (M + 24, y + 67),
               "① 上电并复位后，车辆保持不动约 2.3 秒；② 程序等待 1 秒并采样 Z 轴 256 次；\n"
               "③ 屏幕 HDG 从 0°附近开始；④ 静止时 HDG 不应持续单向增长。",
               font(25), width=PAGE_W - 2 * M - 48, leading=8)
    path = tmp / "page-1.png"
    img.save(path, "PNG")
    paths.append(path)

    # Page 2: gyro control policy and tuning procedure.
    img, d = page_base(2, "SMART CAR / HEADING CONTROL")
    d.text((M, 150), "陀螺仪的作用与调节方法", font=font(55, True), fill=NAVY)
    d.text((M, 227), "用来保持航向；不替代灰度在弯道中的循迹控制", font=font(28), fill=MUTED)

    y = heading(d, 307, "当前控制边界", "1")
    y = bullet_list(d, y, [
        "弯道中：以 8 路灰度传感器为主。半圆内车头自然改变，不应把这段角度变化误判为偏航后强拉回。",
        "直线中：当灰度稳定回到 X4/X5 或确认离开弯道后，才使用陀螺仪保持目标航向。",
        "全白/丢线：必须先由赛道状态机判断“已出弯并且进入直线恢复阶段”，再以目标航向做强制摆正；不能只因一次全白就立即转向。",
        "灰度偏在 X6/X7 或 X2/X3 时，优先检查灰度映射与曲线目标；它不是陀螺仪需要回零的信号。",
    ], size=27, gap=15)

    y += 8
    y = heading(d, y, "验证与微调顺序", "2")
    box(d, (M, y, PAGE_W - M, y + 454), PALE_BLUE)
    steps = [
        ("1", "先静止：上电复位后不碰车，确认 HDG 停在接近 0°。"),
        ("2", "手动转 90° / 180° / 360°，记录 OLED 的 HDG。"),
        ("3", "当前 360° 为 356.6°，误差约 −3.4°（约 0.94%），属于可接受范围，先维持比例 1.00。"),
        ("4", "若静止仍慢慢增长：先检查车辆是否振动；再把静止阈值从 120 小幅提高到 130–140。"),
        ("5", "若整圈误差持续超过约 7°：才调整比例；每次只改 1–2%，重新完整转一圈验证。"),
    ]
    sy = y + 26
    for no, desc in steps:
        d.ellipse((M + 26, sy, M + 62, sy + 36), fill=TEAL)
        d.text((M + 37, sy + 4), no, font=font(21, True), fill=WHITE)
        sy = text_block(d, (M + 84, sy - 2), desc, font(25), width=PAGE_W - 2 * M - 114, leading=7) + 12
    y += 495

    y = heading(d, y, "代码位置与实时观测", "3")
    y = simple_table(d, y, ["文件", "关键对象", "说明"], [
        ("oled_test.c", "g_oled_test_heading_mdeg", "OLED 测试版的实时航向（毫度）"),
        ("oled_test.c", "g_oled_test_gyro_z_bias", "当前 Z 轴零偏（raw）"),
        ("oled_test.c", "静止阈值常量（120）", "OLED_TEST_STATIONARY_RAW_THRESHOLD"),
        ("line_follow_test.c", "GYRO_HEADING_SCALE_NUM / DEN", "循迹程序使用的航向比例参数"),
    ], [300, 450, 568], row_height=68, text_size=22)

    y += 35
    box(d, (M, y, PAGE_W - M, y + 285), PALE_GREEN)
    d.text((M + 24, y + 21), "本轮测试记录", font=font(29, True), fill=TEAL)
    text_block(d, (M + 24, y + 72),
               "早期误差主要来自把刷新周期固定当作 10 ms。现在按 SysTick 的实际时间积分，并在静止时关闭航向积分、\n"
               "缓慢更新 bias。最新读数：整圈 356.6°；转回初始方向 1.8°。",
               font(25), width=PAGE_W - 2 * M - 48, leading=8)
    d.text((M + 24, y + 210), "结论：先不要放大航向修正；需要先保证“出弯状态机”在正确时刻才启用回正。",
           font=font(24, True), fill=NAVY)
    path = tmp / "page-2.png"
    img.save(path, "PNG")
    paths.append(path)

    # Page 3: OLED wiring, screen fields and test workflow.
    img, d = page_base(3, "SMART CAR / OLED STATUS PANEL")
    d.text((M, 150), "OLED 使用与状态判读", font=font(58, True), fill=NAVY)
    d.text((M, 230), "仅适用于 4 针 SSD1306 I²C OLED（128 × 64）", font=font(28), fill=MUTED)

    y = heading(d, 309, "OLED 接线", "1")
    y = simple_table(d, y, ["OLED 引脚", "TI 板连接", "注意"], [
        ("VCC", "3.3 V", "绝不能接 5 V"),
        ("GND", "GND", "与主板、驱动共地"),
        ("SCL", "PA24 / PINCM54", "软件 I²C 时钟"),
        ("SDA", "PA18 / PINCM40", "软件 I²C 数据"),
    ], [225, 530, 563], row_height=70, text_size=24)
    y += 25
    box(d, (M, y, PAGE_W - M, y + 151), PALE_GREEN)
    d.text((M + 24, y + 20), "接线要点", font=font(28, True), fill=TEAL)
    text_block(d, (M + 24, y + 65),
               "PA18/PA24 不是硬件 I²C 对；当前程序采用开漏软件 I²C。\n"
               "陀螺仪仍使用 PA28/PA31（I²C0），两者互不占用。",
               font(25), width=PAGE_W - 2 * M - 48, leading=8)

    y += 186
    y = heading(d, y, "屏幕字段说明", "2")
    y = simple_table(d, y, ["显示", "含义", "怎样判断"], [
        ("HDG", "车头累计航向角（°）", "静止时应基本不变；整圈约 360°"),
        ("TGT / ERR", "当前目标角 / 航向差", "测试程序中 TGT 为 0°，循迹版由状态机设定"),
        ("GRY", "灰度黑线掩码与线误差", "X1 在左；黑线 raw 低电平"),
        ("PWM", "电机 A / B 当前占空比", "OLED 测试版应为 00/00，不驱动电机"),
        ("IMU", "传感器健康位 / 通信失败数", "健康位为 1；失败数持续增加需查 I²C"),
        ("BIA", "Z 轴零偏 raw 值", "车辆静止时会缓慢学习，属正常"),
        ("ENC", "左右编码器计数", "悬空转轮时应随对应轮变化"),
    ], [184, 504, 630], row_height=64, text_size=21)

    y += 31
    y = heading(d, y, "OLED 测试程序", "3")
    box(d, (M, y, PAGE_W - M, y + 342), PALE)
    d.text((M + 24, y + 21), "构建：make oled-test", font=font(29, True), fill=BLUE)
    text_block(d, (M + 24, y + 73),
               "烧录文件：motor_forward_test/build/oled_test.out。\n"
               "该程序已接入：陀螺仪、8 路灰度、编码器及电机 PWM 读数；但不会调用 Motor_run*，\n"
               "因此用于桌面检查时，车轮应保持停止。OLED 会自动尝试地址 0x3C，再尝试 0x3D。",
               font(25), width=PAGE_W - 2 * M - 48, leading=8)
    d.text((M + 24, y + 263), "若黑屏：先核对 VCC=3.3 V、共地、SCL=PA24、SDA=PA18，再检查是否为 4 针 I²C 型屏。",
           font=font(23, True), fill=NAVY)
    path = tmp / "page-3.png"
    img.save(path, "PNG")
    paths.append(path)
    return paths


def build_docx(paths):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    for index, page in enumerate(paths):
        if index:
            doc.add_section(WD_SECTION.NEW_PAGE)
            section = doc.sections[-1]
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = section.bottom_margin = Inches(0)
            section.left_margin = section.right_margin = Inches(0)
            section.header_distance = section.footer_distance = Inches(0)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.add_run().add_picture(str(page), width=Inches(8.27), height=Inches(11.69))
    doc.save(OUT)


def main():
    with TemporaryDirectory(prefix="smartcar_imu_oled_doc_") as tmp_dir:
        pages = create_pages(Path(tmp_dir))
        build_docx(pages)
    print(OUT)


if __name__ == "__main__":
    main()
