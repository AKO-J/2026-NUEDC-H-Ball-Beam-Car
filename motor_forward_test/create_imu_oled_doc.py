from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("智能车_陀螺仪校准与OLED使用指南_2026-07-27.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
MUTED = "5B6573"
GREEN = "1F6B45"
GOLD = "7A5A00"
RED = "9B1C1C"
# macOS system font with Simplified-Chinese glyphs; use its registered
# PostScript family name so LibreOffice can resolve the East-Asian run font.
FONT = "STHeitiSC-Medium"
BODY_SIZE = 10.5
USABLE_DXA = 9360


def set_run_font(run, size=BODY_SIZE, bold=None, color=INK, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, header=True):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths_dxa):
        col.set(qn("w:w"), str(width))
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[col_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, size=9.5,
                                 bold=(header and row_index == 0),
                                 color=INK)
            if header and row_index == 0:
                shade(cell, LIGHT_BLUE)


def add_text_cell(cell, text, bold=False, color=INK, align=None):
    p = cell.paragraphs[0]
    p.clear()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=bold, color=color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        add_text_cell(cell, value, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            add_text_cell(cell, str(value))
    set_table_geometry(table, widths, header=True)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_label_detail_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    for label, detail in rows:
        cells = table.add_row().cells
        add_text_cell(cells[0], label, bold=True, color=DARK_BLUE)
        add_text_cell(cells[1], detail)
        shade(cells[0], LIGHT_GRAY)
    set_table_geometry(table, [2700, 6660], header=False)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, title, text, color=INK):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, CALLOUT)
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    set_table_geometry(table, [9360], header=False)
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "  ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=color)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        size, color = 16, BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size, color = 13, BLUE
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return p


def add_para(doc, text="", bold_prefix=None, color=INK, after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True, color=color)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p


def apply_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("智能车 H 题｜陀螺仪与 OLED 调试记录")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("2026-07-27  ·  现场调试参考")
    set_run_font(r, size=8.5, color=MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("智能车陀螺仪校准与 OLED 使用指南")
    set_run_font(r, size=24, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("适用平台：MSPM0G3507 TI 板 · ICM-42688-P · SSD1306 128×64 I²C OLED")
    set_run_font(r, size=10.5, color=MUTED)

    rows = [
        ("文档目的", "记录当前已验证的陀螺仪零漂/比例修正与 OLED 诊断方法"),
        ("当前诊断固件", "motor_forward_test/build/oled_test.out"),
        ("安全约束", "OLED 测试只读取状态；电机桥保持释放，PWM 应显示 00/00"),
        ("最终手动验证", "完整一圈 356.6°；回原位 1.8°"),
    ]
    add_label_detail_table(doc, rows)


def add_gyro_section(doc):
    add_heading(doc, "1. 陀螺仪：当前工作方式", 1)
    add_callout(
        doc,
        "当前结论",
        "角度积分已改为实测采样间隔；静止时不积分角度并持续修正零偏。"
        "当前精度足以用于半圆出弯后的航向保持，不用于替代弯中灰度循迹。",
        GREEN,
    )

    add_heading(doc, "1.1 接线与方向", 2)
    add_table(doc, ["信号", "TI 板引脚", "说明"], [
        ("SDA", "PA28（I²C0 SDA，PINCM3）", "ICM-42688 数据线"),
        ("SCL", "PA31（I²C0 SCL，PINCM6）", "ICM-42688 时钟线"),
        ("VCC / GND", "3.3 V / GND", "不得以 5 V 供电"),
        ("CS / AD0", "CS→3.3 V；AD0→GND", "I²C 地址 0x68；程序也会探测 0x69"),
    ], [1700, 3500, 4160])

    add_heading(doc, "1.2 启动校准流程", 2)
    for item in [
        "上电后将小车静止放平；不要在启动后的约 2.3 秒内触碰或旋转小车。",
        "ICM 初始化完成后先热机 1 秒，避免 MEMS 芯片刚上电时的温度漂移进入零点。",
        "连续读取 256 个 Z 轴样本，每个样本间隔 5 ms，取平均值作为 Z 轴初始零偏 B。",
        "使用 SysTick 记录真实采样间隔，而不是假设每次循环固定 10 ms。",
        "若修正后的 Z 轴角速度绝对值不超过 120 原始计数（约 7.3°/s），判为静止：不累计角度，并以 Q8 小数精度慢速更新零偏。",
        "超过静止阈值时才积分车头角度；因此正常车辆转弯会记录，停着时 HDG 不会爬升。",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "1.3 当前关键参数", 2)
    add_table(doc, ["参数", "当前值", "作用"], [
        ("量程换算", "16.4 LSB/(°/s)", "ICM-42688 默认陀螺仪量程"),
        ("比例系数", "1.00", "由 360°手动标定得到；已取消旧的 1.06 放大"),
        ("热机时间", "1000 ms", "降低上电温漂"),
        ("初始零偏样本", "256 × 5 ms", "约 1.28 秒静止平均"),
        ("静止阈值", "|修正后 Z| ≤ 120", "低于约 7.3°/s 时锁角并学习零偏"),
        ("航向单位", "mdeg（毫度）", "OLED 以度显示一位小数；代码内部 90°=90000 mdeg"),
    ], [2100, 2100, 5160])

    add_heading(doc, "1.4 已记录的标定结果", 2)
    add_table(doc, ["阶段", "结果", "结论"], [
        ("旧版固定 10 ms 积分", "90°仅约 68.2°；360°约 241.9°", "循环实际耗时受 OLED/灰度读取影响，固定周期不可用"),
        ("实测 SysTick 周期后", "360°约 367.3°，回原位约 -16.3°", "比例基本正确，但启动温漂/静止零偏仍需压制"),
        ("最终：热机 + Q8 零偏 + 静止锁角", "360°约 356.6°，回原位约 1.8°", "当前配置通过；手动转角误差在可接受范围"),
    ], [2400, 3000, 3960])

    add_heading(doc, "1.5 运行策略（用于循迹程序）", 2)
    add_bullet(doc, "黑线可见、处在半圆弯道时：灰度传感器决定转向；不要用陀螺仪把 X6/X7 的正常预瞄状态强拉回 X4/X5。")
    add_bullet(doc, "弯道后连续全白确认完成时：才允许陀螺仪按预期 0°或 ±180°切线做直线航向保持。")
    add_bullet(doc, "直线航向保持使用已人工校正的左右基础 PWM；陀螺仪只加差速修正，不能覆盖灰度正在确认的弯道。")


def add_oled_section(doc):
    add_heading(doc, "2. OLED：接线、显示与诊断", 1)
    add_callout(
        doc,
        "重要",
        "PA18/PA24 不是同一组硬件 I²C 引脚。当前程序使用软件 I²C：PA24 为 SCL，PA18 为 SDA。",
        GOLD,
    )

    add_heading(doc, "2.1 OLED 接线（四针 I²C SSD1306）", 2)
    add_table(doc, ["OLED 引脚", "连接到 TI 板", "注意事项"], [
        ("VCC", "3.3 V", "不要接 5 V"),
        ("GND", "GND", "必须与 TI 板共地"),
        ("SCL", "PA24（PINCM54）", "软件 I²C 时钟"),
        ("SDA", "PA18（PINCM40）", "软件 I²C 数据"),
    ], [1700, 3200, 4460])
    add_para(doc, "适用模块：GND / VCC / SCL / SDA 四针 I²C SSD1306，默认 128×64。程序会依次尝试地址 0x3C 与 0x3D。", color=MUTED)
    add_para(doc, "不适用：带 CS、DC、RES 等额外引脚的 SPI OLED；此类模块需要改用 SPI 驱动。", bold_prefix="不适用：", color=RED)

    add_heading(doc, "2.2 屏幕字段说明", 2)
    add_table(doc, ["字段", "来源", "正常现象 / 用途"], [
        ("HDG", "ICM 实时积分角度", "静止时应基本稳定；转动小车时变化"),
        ("TGT / ERR", "当前目标角度 / HDG−TGT", "诊断程序目标为 0°；循迹版可显示出弯目标"),
        ("S=xx", "8 路灰度黑线掩码（十六进制）", "0x18 表示 X4、X5 在黑线；bit0=X1，bit7=X8"),
        ("E", "灰度加权横向偏差", "负数靠 X1 左侧；正数靠 X8 右侧"),
        ("PWM", "TB6612 A/B 当前输出", "诊断程序必须为 00/00；循迹时显示当前指令"),
        ("IMU / F", "姿态传感器在线 / 连续失败次数", "IMU1F00 表示当前读取正常"),
        ("B", "当前 Z 轴零偏", "静止时可缓慢变化，属于温漂自校正"),
        ("ENC", "左右编码器累计计数", "手转对应车轮时应变化；用于确认编码器接线"),
        ("P", "诊断相位标识", "P1=ICM 已就绪；P0=正在重试或初始化失败"),
    ], [1500, 3180, 4680])

    add_heading(doc, "2.3 构建与烧录 OLED 诊断程序", 2)
    add_para(doc, "项目目录：仓库内 motor_forward_test/", bold_prefix="项目目录：")
    add_table(doc, ["动作", "命令 / 目标"], [
        ("构建", "make oled-test"),
        ("输出固件", "build/oled_test.out"),
        ("烧录性质", "只读诊断：初始化后 Motor_stopAll()；不调用任何 Motor_run* 函数"),
    ], [2100, 7260])
    add_para(doc, "烧录后请先静止约 3 秒。此固件会读取：OLED、ICM42688、8 路灰度、编码器与电机 PWM 状态；不会让车轮自行转动。", color=GREEN)

    add_heading(doc, "2.4 快速验收顺序", 2)
    for item in [
        "确认屏幕显示而非黑屏；若黑屏，先核对 VCC=3.3 V、GND、SCL=PA24、SDA=PA18。",
        "静止观察 HDG 约 20 秒：应不再单向持续爬升；IMU 应显示 1。",
        "手动左右转动小车：HDG 应随方向变化；完成一圈后读数应接近 360°。",
        "将黑线依次移至 X1、X4/X5、X8：S 与 E 应分别表现为左、居中、右。",
        "手转左右驱动轮：ENC 两个通道应能变化；PWM 在诊断程序中继续保持 00/00。",
    ]:
        add_numbered(doc, item)


def add_troubleshooting(doc):
    add_heading(doc, "3. 故障排查与维护", 1)
    add_table(doc, ["现象", "优先检查", "处理方式"], [
        ("OLED 黑屏", "模块是否为 I²C 四针；供电与地；SCL/SDA 是否对调", "使用 3.3 V；确认 PA24=SCL、PA18=SDA；检查 0x3C/0x3D 模块地址"),
        ("HDG 不动", "IMU 是否显示 1；是否一直处在静止阈值内", "快速转动超过静止阈值；检查 PA28/PA31 与 ICM 供电"),
        ("HDG 静止爬升", "是否在启动热机/取零阶段移动了小车", "复位后静止约 3 秒；仍异常则记录 B、IMU/F、HDG 供复核"),
        ("360°比例偏差", "先做完整一圈再评估，避免单次 90°手动角度误差", "比例校正 = 旧比例 × 360 / 实测一圈角度；一次只改一项"),
        ("ENC 不变", "编码器供电、E1A/E1B/E2A/E2B 信号、手转是否真正带动轴", "先用 OLED 诊断排线，再进入电机闭环调速"),
    ], [1700, 3300, 4360])

    add_heading(doc, "4. 当前文件对应关系", 1)
    add_table(doc, ["文件", "职责"], [
        ("oled_test.c", "安全的 OLED + ICM + 灰度 + 编码器只读诊断程序"),
        ("ssd1306_oled.c / .h", "PA24/PA18 软件 I²C SSD1306 驱动及屏幕排版"),
        ("icm42688.c / .h", "ICM42688 I²C0（PA28/PA31）初始化与采样"),
        ("line_follow_test.c", "正式循迹控制；与诊断程序共享当前陀螺仪比例 1.00"),
        ("Makefile", "目标命令：make oled-test；正式循迹目标为 make line-follow"),
    ], [3300, 6060])
    add_para(doc, "文档范围：本指南记录 2026-07-27 已验证的诊断与校准状态。变更接线、陀螺仪安装方向、量程或采样逻辑后，必须重新完成完整一圈与回原位测试。", color=MUTED, italic=True)


def build():
    doc = Document()
    apply_styles(doc)
    add_title(doc)
    add_gyro_section(doc)
    add_oled_section(doc)
    add_troubleshooting(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
