from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).parent
OUT = ROOT / "智能车_正方形循迹_最新状态.docx"
WORK = Path("/private/tmp/smartcar_status_doc_pages")

W, H = 1300, 1770
NAVY = "#0B2545"
BLUE = "#2E74B5"
MID_BLUE = "#DCE9F5"
PALE_BLUE = "#F1F6FA"
PALE_GRAY = "#F4F6F8"
GRAY = "#596675"
WHITE = "#FFFFFF"
BLACK = "#102A43"

FONT = "/System/Library/Fonts/STHeiti Light.ttc"
FONT_BOLD = "/System/Library/Fonts/STHeiti Medium.ttc"


def f(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT, size, index=0)


def tw(draw, text, font):
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0]


def wrap(draw, text, font, width):
    """Character-aware wrapping for Chinese mixed with paths and identifiers."""
    lines, current = [], ""
    for part in text.split("\n"):
        current = ""
        for char in part:
            candidate = current + char
            if current and tw(draw, candidate, font) > width:
                lines.append(current)
                current = char
            else:
                current = candidate
        lines.append(current or " ")
    return lines or [" "]


def text_block(draw, xy, text, font, fill, width, line_gap=8):
    x, y = xy
    lines = wrap(draw, text, font, width)
    line_h = font.size + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h
    return y


def page_base(page_no, title, subtitle):
    image = Image.new("RGB", (W, H), WHITE)
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, W, 15), fill=NAVY)
    draw.text((56, 42), title, font=f(38, True), fill=NAVY)
    draw.text((58, 95), subtitle, font=f(19), fill=GRAY)
    draw.line((56, 132, W - 56, 132), fill=BLUE, width=3)
    footer = f"智能车循迹 · 最新精简状态 · 2026-07-26 · {page_no}/3"
    draw.text((56, H - 45), footer, font=f(16), fill=GRAY)
    return image, draw, 165


def heading(draw, y, text):
    draw.text((56, y), text, font=f(26, True), fill=BLUE)
    return y + 44


def callout(draw, y, label, text):
    x = 56
    width = W - 112
    label_font = f(20, True)
    text_font = f(20)
    lines = wrap(draw, text, text_font, width - 42)
    height = 52 + len(lines) * 28
    draw.rounded_rectangle((x, y, x + width, y + height), radius=14, fill=PALE_BLUE)
    draw.text((x + 20, y + 16), label, font=label_font, fill=NAVY)
    text_y = y + 16
    label_w = tw(draw, label + "  ", label_font)
    # Keep the first sentence visually compact; body continues on the next line.
    first = lines[0]
    draw.text((x + 20 + label_w, text_y), first, font=text_font, fill=BLACK)
    for index, line in enumerate(lines[1:], 1):
        draw.text((x + 20, text_y + index * 28), line, font=text_font, fill=BLACK)
    return y + height + 20


def table(draw, y, headers, rows, widths, size=18):
    x = 56
    total = sum(widths)
    header_font = f(size, True)
    body_font = f(size)
    line_h = size + 8

    def row_height(values, font):
        counts = [len(wrap(draw, value, font, width - 20)) for value, width in zip(values, widths)]
        return max(34, max(counts) * line_h + 20)

    h = row_height(headers, header_font)
    draw.rounded_rectangle((x, y, x + total, y + h), radius=10, fill=MID_BLUE)
    cx = x
    for value, width in zip(headers, widths):
        text_block(draw, (cx + 10, y + 10), value, header_font, NAVY, width - 20, 8)
        cx += width
    y += h

    for index, values in enumerate(rows):
        h = row_height(values, body_font)
        fill = WHITE if index % 2 == 0 else PALE_GRAY
        draw.rectangle((x, y, x + total, y + h), fill=fill)
        cx = x
        for col, (value, width) in enumerate(zip(values, widths)):
            draw.rectangle((cx, y, cx + width, y + h), outline="#D5DEE7", width=1)
            text_block(
                draw,
                (cx + 10, y + 10),
                value,
                f(size, bold=(col == 0)),
                BLACK,
                width - 20,
                8,
            )
            cx += width
        y += h
    return y + 20


def bullets(draw, y, items, size=20):
    font = f(size)
    for item in items:
        lines = wrap(draw, item, font, W - 150)
        draw.ellipse((68, y + 10, 77, y + 19), fill=BLUE)
        for index, line in enumerate(lines):
            draw.text((92, y + index * (size + 8)), line, font=font, fill=BLACK)
        y += len(lines) * (size + 8) + 11
    return y


def numbered(draw, y, items):
    font = f(19)
    for index, item in enumerate(items, 1):
        lines = wrap(draw, item, font, W - 170)
        draw.text((66, y), f"{index}.", font=f(19, True), fill=BLUE)
        for line_no, line in enumerate(lines):
            draw.text((112, y + line_no * 27), line, font=font, fill=BLACK)
        y += len(lines) * 27 + 12
    return y


def make_pages():
    WORK.mkdir(parents=True, exist_ok=True)

    # Page 1: factual setup and current state machine.
    image, draw, y = page_base(1, "智能车 · 正方形循迹", "当前技术状态与下一步地面验证")
    y = callout(
        draw,
        y,
        "当前结论",
        "软件已具备“角点触发 → 主动刹车 → ICM42688 累计 90° → 重新找线”的完整流程。单个直角的地面稳定性尚未验收。",
    )
    y = heading(draw, y, "1. 已确认的硬件与标定")
    y = table(
        draw,
        y,
        ("项目", "当前确认", "备注"),
        [
            ("主控与驱动", "MSPM0G3507 + TB6612", "电机 A 是物理左轮，B 是物理右轮。"),
            ("灰度模块", "8 路 Yahboom", "车头向前时 X1 在左；黑线原始电平为 0。"),
            ("IMU", "ICM42688", "I2C 通信已验证；Z 轴用于航向累计。"),
            ("几何", "轮距 112 mm；前伸 152 mm", "X1–X8 跨度 79 mm；离地 19 mm。"),
            ("赛道", "胶带 19 mm；最小弯半径 400 mm", "当前目标：稳定通过 90° 方形角。"),
        ],
        (235, 405, 548),
        17,
    )
    y = heading(draw, y, "2. 当前有效控制逻辑")
    y = callout(
        draw,
        y,
        "角点条件",
        "先沿窄线（≤3 路黑）运行 24 个周期（约 96 ms）；再出现同侧 ≥4 路黑的窄→宽跃升才转弯。静态放在宽黑区不会立即转。",
    )
    y = table(
        draw,
        y,
        ("状态", "核心动作", "退出条件"),
        [
            ("0 FOLLOW", "4 ms 周期，单次 8 路扫描。", "有效窄→宽角点边沿。"),
            ("1 BRAKE", "TB6612 双轮主动短刹 24 ms。", "进入 IMU 转弯。"),
            ("2 TURN", "单轮 PIVOT，累计航向。", "达到 90,000 mdeg。"),
            ("3 REACQUIRE", "18% 低速前进找新边。", "连续 2 帧窄线；最多 600 ms。"),
            ("4 FAULT", "安全停车。", "IMU 读错、转弯或找线超时。"),
        ],
        (210, 470, 508),
        17,
    )
    image.save(WORK / "page1.png")

    # Page 2: exact turning command and the parameters worth touching.
    image, draw, y = page_base(2, "现行控制逻辑与参数", "只保留当前有效路径；历史反复方案已移除")
    y = heading(draw, y, "3. 转弯方向与刹车")
    y = bullets(
        draw,
        y,
        [
            "右侧支路（X5–X8）：左轮驱动、右轮 0%，形成右转。",
            "左侧支路（X1–X4）：左轮 0%、右轮驱动，形成左转。",
            "IMU 累计到 72° 后降速；达到 90° 后主动短刹 32 ms，再低速找下一条线。",
            "主动刹车为 TB6612 两个方向输入同时有效，不是普通的松开/滑行停止。",
        ],
        21,
    )
    y = heading(draw, y, "4. 当前关键参数")
    y = table(
        draw,
        y,
        ("参数", "当前值", "用途"),
        [
            ("LINE_FOLLOW_CONTROL_SLICE_MS", "4 ms", "普通循迹和角点读取周期。"),
            ("SQUARE_PIVOT_ARM_NARROW_SLICES", "24", "约 96 ms 窄线预置，防静态误转。"),
            ("SQUARE_BRAKE_MS", "24 ms", "角点触发后的主动短刹。"),
            ("SQUARE_TURN_TARGET_MDEG", "90,000", "IMU 90° 转弯结束条件。"),
            ("SQUARE_TURN_SLOWDOWN_MDEG", "72,000", "最后 18° 的低速收角起点。"),
            ("右侧支路左轮 PWM", "41% → 18%", "右转主驱动轮。"),
            ("左侧支路右轮 PWM", "20% 起转 → 10% 保持", "左转主驱动轮。"),
            ("SQUARE_TURN_FINISH_BRAKE_MS", "32 ms", "IMU 到 90° 后抑制惯性。"),
            ("SQUARE_REACQUIRE_DUTY / MAX_MS", "18% / 600 ms", "转完后重新找到线。"),
        ],
        (475, 260, 453),
        16,
    )
    y = heading(draw, y, "可调优先级")
    bullets(
        draw,
        y,
        [
            "角度不足或超过：先改 SQUARE_TURN_TARGET_MDEG。",
            "转完后找不到线：改 SQUARE_REACQUIRE_DUTY；一次只改一项。",
            "不要先提高急停或转弯 PWM；先完成单角闭环验证。",
        ],
        20,
    )
    image.save(WORK / "page2.png")

    # Page 3: operational handoff.
    image, draw, y = page_base(3, "构建、验证与调试", "按下面顺序进行下一次地面试验")
    y = heading(draw, y, "5. 代码、构建与烧录")
    y = bullets(
        draw,
        y,
        [
            "主程序：motor_forward_test/line_follow_test.c。",
            "电机控制：motor_forward_test/motor_driver.c；Motor_brakeAllFor() 是最新主动刹车接口。",
            "灰度与 IMU：gray_sensor.c、icm42688.c；主程序采用单次 8 路扫描提升角点采样速度。",
            "构建：在 motor_forward_test 目录执行 make line-follow。烧录产物：build/line_follow_test.out。",
            "烧录后上电前约 4 秒不要移动小车，以完成 IMU 零偏标定。",
        ],
        19,
    )
    y = heading(draw, y, "6. 下一步地面验证")
    y = numbered(
        draw,
        y,
        [
            "只布置“直线 + 一个右侧 90° 角”，从普通窄线段开始；不要把车直接放在宽黑区域。",
            "上电后保持静止约 4 秒，确认直线前进正常、不会静态误转。",
            "观察角点序列：主动短刹 → 左轮驱动、右轮停 → 接近 90° 降速并短刹 → 低速找新边。",
            "若出现问题，结合 Watch 判断：角度问题先调 TARGET；找线问题先调 REACQUIRE；每次只改一项。",
            "连续通过单角 5 次后，再测试完整正方形赛道。",
        ],
    )
    y = heading(draw, y, "7. CCS Watch（最小集合）")
    y = table(
        draw,
        y,
        ("变量", "正常应看到"),
        [
            ("g_square_state", "0 → 1 → 2 → 3 → 0；变为 4 代表安全故障停车。"),
            ("g_square_turn_angle_mdeg", "转弯结束接近 90,000。"),
            ("g_square_turn_driven_wheel_is_left", "右侧支路为 1；左侧支路为 0。"),
            ("g_line_applied_left_duty / right_duty", "TURN 时核对实际 PWM 是否符合物理轮命令。"),
            ("g_square_imu_error / g_square_turn_timed_out", "正常均为 0。"),
        ],
        (475, 713),
        15,
    )
    y = heading(draw, y, "已废弃：不要恢复")
    bullets(
        draw,
        y,
        [
            "“黑线数降到一半以下即结束转弯”——会在直角处过早结束。",
            "内外轮速度比、编码器弯道比 2.17、SQUARE_TURN_REVERSE_STEERING 反转映射。",
            "右轮持续 1% PWM——通常不足以让电机从静止起转。",
        ],
        18,
    )
    image.save(WORK / "page3.png")


def build_docx():
    make_pages()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    for index in range(1, 4):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Inches(0)
        paragraph.paragraph_format.space_after = Inches(0)
        run = paragraph.add_run()
        run.add_picture(str(WORK / f"page{index}.png"), width=Inches(6.5), height=Inches(8.85))
        if index < 3:
            run.add_break(WD_BREAK.PAGE)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
